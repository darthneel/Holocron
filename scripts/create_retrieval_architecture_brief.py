from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "docs" / "retrieval-architecture-options.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "6B7280"
HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(round(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_fixed(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def para(doc, text="", *, style=None, before=None, after=None, line=1.25, color=INK, size=11, bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    pf = p.paragraph_format
    pf.line_spacing = line
    if before is not None:
        pf.space_before = Pt(before)
    if after is not None:
        pf.space_after = Pt(after)
    r = p.add_run(text)
    set_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def bullet(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.375)
    pf.first_line_indent = Inches(-0.188)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.25
    r = p.add_run("•  " + text)
    set_font(r)
    return p


def add_heading(doc, text, level=1):
    values = {1: (16, BLUE, 18, 10), 2: (13, BLUE, 14, 7), 3: (12, DARK_BLUE, 10, 5)}
    size, color, before, after = values[level]
    return para(doc, text, before=before, after=after, size=size, color=color, bold=True)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_fixed(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, text in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_font(run, size=10, color=DARK_BLUE, bold=True)
    for data in rows:
        cells = table.add_row().cells
        for i, text in enumerate(data):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            set_font(run, size=10)
    set_table_fixed(table, widths)
    return table


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_fixed(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + " ")
    set_font(r, size=10.5, color=DARK_BLUE, bold=True)
    r = p.add_run(text)
    set_font(r, size=10.5)
    para(doc, "", after=0, size=1)


def set_header_footer(section):
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("HOLOCRON  /  RETRIEVAL ARCHITECTURE")
    set_font(run, size=8.5, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)
    run = footer.add_run("  |  Internal reference")
    set_font(run, size=8.5, color=MUTED)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)
    set_header_footer(sec)

    # Compact reference guide preset + memo-masthead header pattern.
    title = para(doc, "Retrieval Architecture Options", before=0, after=4, size=23, color="000000", bold=True)
    subtitle = para(doc, "A practical path beyond a basic RAG pipeline", after=12, size=14, color="434343")
    meta = para(doc, "Purpose: guide the next stages of Holocron’s briefing retrieval system", after=14, size=10, color=MUTED)

    add_callout(doc, "Bottom line.", "Keep PostgreSQL and pgvector as the retrieval foundation. The highest-leverage next step is a hybrid, routed retrieval system: use structured queries when the question is structured; otherwise combine lexical, semantic, recency, and importance signals, then rerank a small, workspace-authorized candidate set.")

    add_heading(doc, "Why a basic RAG system is not the end state")
    para(doc, "A straightforward vector search is an important first improvement because it can recover relevant material that is not among the newest linked events. But a briefing system needs more than topical similarity: it also needs factual precision, relationship awareness, time sensitivity, and a defensible reason for every included claim.")

    add_heading(doc, "Six capabilities worth adding", 2)
    rows = [
        ("1. Hybrid retrieval", "Combine lexical/full-text, vector similarity, recency, importance, and relationship signals. Fuse candidates, then rerank."),
        ("2. Retrieval routing", "Choose the retrieval method from the question: structured SQL for structured facts; semantic or hybrid search for narrative context."),
        ("3. Reranking", "Apply a more precise relevance model to the small candidate set after first-pass retrieval."),
        ("4. Importance-aware memory", "Let operators pin durable commitments, sensitive history, decisions, and strategic priorities so they do not decay with recency."),
        ("5. Event/decision memory", "Represent decisive facts as first-class records, not only as text chunks embedded in notes."),
        ("6. Contextual enrichment", "Preserve document metadata—participants, date, topic, source—alongside each chunk to make retrieval and citations more reliable."),
    ]
    add_table(doc, ["Capability", "What it contributes"], rows, [1.875, 4.625])

    add_heading(doc, "Retrieval should follow the question", 2)
    route_rows = [
        ("Who is attending?", "Relational/calendar query"),
        ("What did we promise this person?", "Linked history plus semantic retrieval"),
        ("What did we discuss about permitting?", "Hybrid semantic and keyword retrieval"),
        ("When is the next meeting?", "Structured SQL/calendar query"),
        ("Which commitments are overdue?", "Workflow/status query"),
    ]
    add_table(doc, ["Need", "Best retrieval path"], route_rows, [2.45, 4.05])

    add_heading(doc, "The recommended architecture")
    para(doc, "This keeps the system inspectable while improving recall and reducing unnecessary context sent to the model.")
    add_table(doc, ["Stage", "Responsibility"], [
        ("1. Structured facts + relationship graph", "Meetings, people, workspaces, ownership, statuses, dates, and direct links."),
        ("2. Hybrid candidate retrieval", "Within the authorized workspace, collect lexical, vector, recency, and importance candidates."),
        ("3. Rerank + permission checks", "Select the few most useful items after applying workspace and access constraints."),
        ("4. Cited context manifest", "Pass a compact, traceable set of source snippets and structured facts to the briefing generator."),
        ("5. Briefing generation", "Generate claims only from the selected manifest; record provenance and token use."),
    ], [2.3, 4.2])

    add_heading(doc, "What this buys you", 2)
    for item in [
        "Better recall than ‘last few linked meetings’ alone, without treating every old note as equally important.",
        "Lower input-token cost because the model receives a smaller, higher-signal context manifest.",
        "More useful citations: each claim can point back to a meeting, note, decision, or structured record.",
        "A clean evaluation loop: compare useful cited claims ÷ input tokens across linked-recency and semantic/hybrid paths.",
        "Workspace scoping remains a hard retrieval boundary, not merely a prompt instruction.",
    ]:
        bullet(doc, item)

    add_heading(doc, "Represent durable memory explicitly", 2)
    para(doc, "Do not leave consequential facts trapped in prose. A durable memory record can preserve the fact, its owner, date, source, confidence, expiry/review policy, and access scope.")
    add_table(doc, ["Field", "Example"], [
        ("Decision", "Mayor committed to creating a permitting liaison."),
        ("Owner", "Operations lead"),
        ("Date", "March 9"),
        ("Source", "Darius call"),
        ("Review policy", "Pinned; annual review"),
    ], [1.875, 4.625])

    add_heading(doc, "What not to prioritize yet")
    for item in [
        "A separate vector database. Postgres/Neon already keeps authorization, provenance, and operational data close together.",
        "A graph database. The current relational model can represent the relationship graph needed at this stage.",
        "Fine-tuning. It does not solve missing, stale, or incorrectly retrieved context.",
        "Autonomous retrieval agents with authority over access. Keep retrieval constrained, logged, and permission-filtered.",
    ]:
        bullet(doc, item)

    add_heading(doc, "Practical sequence", 2)
    for item in [
        "Keep the current linked-recency path as the control, and retain the semantic path for side-by-side evaluation.",
        "Add full-text scoring and reciprocal-rank fusion to build the hybrid candidate set.",
        "Introduce a reranker only after measuring cases where first-pass retrieval is insufficient.",
        "Add explicit decision and commitment records for high-value facts.",
        "Use the existing comparison metric to decide whether each addition improves briefing quality per input token.",
    ]:
        bullet(doc, item)

    add_heading(doc, "Reference", 2)
    p = para(doc, "pgvector documentation: vector search in Postgres, including hybrid search patterns, reciprocal rank fusion, and reranking guidance.", after=2, size=9.5, color=MUTED)
    p.add_run(" ")
    link = p.add_run("https://github.com/pgvector/pgvector")
    set_font(link, size=9.5, color=BLUE)
    OUT.unlink(missing_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
